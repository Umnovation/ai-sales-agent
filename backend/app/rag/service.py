"""RAG service — document upload, chunking, embedding, retrieval."""
from __future__ import annotations

import os
from pathlib import Path

import structlog
from fastapi import HTTPException, UploadFile
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.rag.embedder import OpenAIEmbedder
from app.rag.models import Document, DocumentChunk

logger: structlog.stdlib.BoundLogger = structlog.get_logger()

CHUNK_SIZE: int = 500  # characters per chunk
CHUNK_OVERLAP: int = 50

ALLOWED_EXTENSIONS: set[str] = {".docx", ".pdf", ".txt"}
ALLOWED_MIME_TYPES: set[str] = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


async def upload_document(
    db: AsyncSession,
    file: UploadFile,
) -> Document:
    """Upload, parse, chunk, embed, and store a document."""
    if file.filename is None:
        raise HTTPException(status_code=400, detail="Filename is required")

    ext: str = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{ext}' not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    content_bytes: bytes = await file.read()
    file_size: int = len(content_bytes)

    if file_size > settings.max_upload_size_bytes:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Max: {settings.max_upload_size_mb}MB",
        )

    # Parse content
    text_content: str = _parse_file(content_bytes, ext)

    # Chunk
    chunks: list[str] = _chunk_text(text_content)

    # Embed — get API key from CompanySettings
    from app.settings.models import CompanySettings

    cs_result = await db.execute(select(CompanySettings).limit(1))
    cs: CompanySettings | None = cs_result.scalar_one_or_none()
    if cs is None or not cs.ai_api_key:
        raise HTTPException(status_code=400, detail="AI API key not configured in Settings.")
    embedder = OpenAIEmbedder(api_key=cs.ai_api_key, model=cs.ai_embedding_model)
    embeddings: list[list[float]] = await embedder.embed(chunks)

    # Save to DB
    document = Document(
        filename=file.filename,
        file_type=ext.lstrip("."),
        file_size=file_size,
        chunk_count=len(chunks),
    )
    db.add(document)
    await db.flush()

    for idx, (chunk_text, embedding) in enumerate(zip(chunks, embeddings, strict=True)):
        chunk = DocumentChunk(
            document_id=document.id,
            content=chunk_text,
            embedding=embedding,
            chunk_index=idx,
        )
        db.add(chunk)

    await db.commit()
    await db.refresh(document)

    # Save file to disk
    upload_dir: Path = Path(settings.upload_dir)
    upload_dir.mkdir(parents=True, exist_ok=True)
    file_path: Path = upload_dir / f"{document.id}_{file.filename}"
    with open(file_path, "wb") as f:
        f.write(content_bytes)

    logger.info(
        "document_uploaded",
        document_id=document.id,
        filename=file.filename,
        chunks=len(chunks),
    )

    return document


async def list_documents(db: AsyncSession) -> list[Document]:
    result = await db.execute(select(Document).order_by(Document.created_at.desc()))
    return list(result.scalars().all())


async def delete_document(db: AsyncSession, document_id: int) -> None:
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    document: Document | None = result.scalar_one_or_none()
    if document is None:
        raise HTTPException(status_code=404, detail="Document not found")

    # Delete file from disk
    upload_dir: Path = Path(settings.upload_dir)
    for file_path in upload_dir.glob(f"{document.id}_*"):
        os.remove(file_path)

    await db.delete(document)
    await db.commit()

    logger.info("document_deleted", document_id=document_id)


async def retrieve_relevant_chunks(
    db: AsyncSession,
    ai_provider: object,  # unused, embedder is separate
    query: str,
    limit: int = 3,
) -> list[str]:
    """Retrieve most relevant document chunks for a query."""
    from app.settings.models import CompanySettings

    cs_result = await db.execute(select(CompanySettings).limit(1))
    cs: CompanySettings | None = cs_result.scalar_one_or_none()
    if cs is None or not cs.ai_api_key:
        return []  # No key = no RAG, silently skip
    embedder = OpenAIEmbedder(api_key=cs.ai_api_key, model=cs.ai_embedding_model)
    query_embeddings: list[list[float]] = await embedder.embed([query])
    query_embedding: list[float] = query_embeddings[0]

    # pgvector cosine similarity search
    result = await db.execute(
        text(
            """
            SELECT content
            FROM document_chunks
            ORDER BY embedding <=> :embedding
            LIMIT :limit
            """
        ).bindparams(
            embedding=str(query_embedding),
            limit=limit,
        )
    )

    rows = result.fetchall()
    return [str(row[0]) for row in rows]


def _parse_file(content: bytes, extension: str) -> str:
    """Parse file content to plain text."""
    if extension == ".txt":
        return content.decode("utf-8", errors="replace")
    elif extension == ".docx":
        import docx
        from io import BytesIO

        doc = docx.Document(BytesIO(content))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    elif extension == ".pdf":
        from io import BytesIO

        from PyPDF2 import PdfReader

        reader = PdfReader(BytesIO(content))
        return "\n".join(
            page.extract_text() or "" for page in reader.pages
        )
    else:
        msg: str = f"Unsupported file type: {extension}"
        raise ValueError(msg)


def _chunk_text(text_content: str) -> list[str]:
    """Split text into overlapping chunks."""
    if not text_content.strip():
        return []

    chunks: list[str] = []
    start: int = 0

    while start < len(text_content):
        end: int = start + CHUNK_SIZE
        chunk: str = text_content[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - CHUNK_OVERLAP

    return chunks
